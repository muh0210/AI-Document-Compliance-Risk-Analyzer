"""
Document Text Extraction Module
Handles PDF, DOCX, and TXT file ingestion.
Compatible with both Flask FileStorage and Streamlit UploadedFile.
"""

import os
import io
import re


def extract_text(file_obj):
    """
    Extract text from an uploaded file.

    Args:
        file_obj: Flask FileStorage or Streamlit UploadedFile (any object
                  with a .read() method and .filename or .name attribute).

    Returns:
        dict with keys: text, filename, file_type, page_count, word_count, char_count
    """
    # Resolve filename across frameworks
    filename = getattr(file_obj, "filename", None) or getattr(file_obj, "name", "document.txt")
    ext = os.path.splitext(filename)[1].lower()

    if ext not in (".pdf", ".docx", ".doc", ".txt"):
        raise ValueError(f"Unsupported file format: {ext}. Supported: PDF, DOCX, TXT")

    raw = file_obj.read()
    # Reset stream pointer so the file can be re-read by callers if needed
    if hasattr(file_obj, "seek"):
        file_obj.seek(0)

    if ext == ".pdf":
        text, pages = _pdf(raw)
    elif ext in (".docx", ".doc"):
        text, pages = _docx(raw)
    else:
        text, pages = _txt(raw)

    text = _normalise(text)

    if not text or len(text.strip()) < 10:
        raise ValueError("Could not extract meaningful text from the document.")

    words = text.split()
    return {
        "text": text,
        "filename": filename,
        "file_type": ext.lstrip(".").upper(),
        "page_count": pages,
        "word_count": len(words),
        "char_count": len(text),
    }


# ── Private helpers ──────────────────────────────────────


def _pdf(raw: bytes):
    from PyPDF2 import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    parts = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            parts.append(t.strip())
    if not parts:
        raise ValueError("PDF appears to be image-based / scanned. OCR is not supported yet.")
    return "\n\n".join(parts), len(reader.pages)


def _docx(raw: bytes):
    from docx import Document

    doc = Document(io.BytesIO(raw))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts), None


def _txt(raw: bytes):
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return raw.decode(enc).strip(), None
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError("Could not decode the text file with any supported encoding.")


def _normalise(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    lines = [ln.strip() for ln in text.split("\n")]
    return "\n".join(lines)
